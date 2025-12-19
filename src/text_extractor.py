"""
テキスト抽出モジュール
PDF・Wordファイルから直接テキストを抽出（OCR不要）
"""
import os
from pathlib import Path
from typing import List, Optional, Callable


def check_docx_available() -> bool:
    """python-docxが利用可能かチェック"""
    try:
        import docx
        return True
    except ImportError:
        return False


class TextExtractor:
    """PDF・Wordからテキストを直接抽出するクラス"""

    def __init__(self):
        pass

    def extract_from_pdf(
        self,
        pdf_path: str,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> List[str]:
        """
        PDFからテキストを直接抽出

        Args:
            pdf_path: PDFファイルパス
            progress_callback: (current, total, status) コールバック

        Returns:
            各ページのテキストのリスト
        """
        import fitz  # PyMuPDF

        results = []
        doc = fitz.open(pdf_path)
        total = len(doc)

        try:
            for page_num in range(total):
                if progress_callback:
                    progress_callback(page_num + 1, total, f"抽出中: {page_num + 1}/{total}ページ")

                page = doc[page_num]
                text = page.get_text()
                results.append(text)

        finally:
            doc.close()

        return results

    def extract_from_docx(
        self,
        docx_path: str,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> List[str]:
        """
        Wordファイル(.docx)からテキストを直接抽出

        Args:
            docx_path: Wordファイルパス
            progress_callback: (current, total, status) コールバック

        Returns:
            段落ごとのテキストのリスト（1要素=1ドキュメント全体）
        """
        if not check_docx_available():
            raise ImportError("python-docxがインストールされていません。\npip install python-docx を実行してください。")

        from docx import Document

        if progress_callback:
            progress_callback(1, 1, "Word文書を読み込み中...")

        doc = Document(docx_path)

        # 全テキストを抽出
        full_text = []

        # 段落を抽出
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # テーブル内のテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append('\t'.join(row_text))

        return ['\n'.join(full_text)]

    def extract_from_doc(
        self,
        doc_path: str,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> List[str]:
        """
        古いWordファイル(.doc)からテキストを抽出
        ※ antiwordまたはLibreOfficeが必要

        Args:
            doc_path: Wordファイルパス
            progress_callback: コールバック

        Returns:
            テキストのリスト
        """
        # .docは複雑なので、.docxへの変換を推奨
        raise NotImplementedError(
            ".doc形式は直接サポートされていません。\n"
            ".docx形式に変換してから使用してください。\n"
            "（Wordで開いて「名前を付けて保存」→「.docx」形式を選択）"
        )

    def save_results(
        self,
        texts: List[str],
        output_path: str,
        page_separator: str = "\n\n--- Page {page} ---\n\n"
    ):
        """抽出結果をテキストファイルに保存"""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            for idx, text in enumerate(texts):
                if idx > 0:
                    f.write(page_separator.format(page=idx + 1))
                f.write(text)

    def extract_to_file(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> str:
        """
        ファイルからテキストを抽出してファイルに保存

        Args:
            input_path: 入力ファイルパス（PDF or DOCX）
            output_path: 出力テキストファイルパス（省略時は自動生成）
            progress_callback: コールバック

        Returns:
            出力ファイルパス
        """
        ext = os.path.splitext(input_path)[1].lower()

        if ext == '.pdf':
            texts = self.extract_from_pdf(input_path, progress_callback)
        elif ext == '.docx':
            texts = self.extract_from_docx(input_path, progress_callback)
        elif ext == '.doc':
            texts = self.extract_from_doc(input_path, progress_callback)
        else:
            raise ValueError(f"サポートされていないファイル形式: {ext}")

        if output_path is None:
            base = os.path.splitext(input_path)[0]
            output_path = f"{base}_text.txt"

        self.save_results(texts, output_path)

        return output_path

    def has_text_content(self, pdf_path: str) -> bool:
        """
        PDFにテキストが含まれているかチェック

        Args:
            pdf_path: PDFファイルパス

        Returns:
            テキストが含まれている場合True
        """
        import fitz

        doc = fitz.open(pdf_path)
        try:
            # 最初の数ページをチェック
            for i in range(min(3, len(doc))):
                page = doc[i]
                text = page.get_text().strip()
                if len(text) > 10:  # 10文字以上あればテキスト付き
                    return True
            return False
        finally:
            doc.close()
